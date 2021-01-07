import  DB
from docx import Document
import addText
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX
from pdf2image import convert_from_path
from docx2pdf import convert
import os
from docx.shared import Pt
import imgkit


def createDir():
    # Directory
    directory = "images"

    # Parent Directory path
    parent_dir = "C:/Users/ASUS/Desktop/"

    # Path
    path = os.path.join(parent_dir, directory)
    os.mkdir(path)

def Summary (docID, threshold=0.5):
    data = DB.get_json(docID)
    dic = addText.convertToText(data)
    document = Document()
    document.add_heading('Summary', 0)
    for i in range(len(dic)):
        if (float(dic[i].get('wight')) > threshold):
            document.add_paragraph(dic[i].get('content').replace("\n",' '))
    #document.add_page_break()
    document.save('demoSummary.docx')
    convertDocxToImage("demoSummary.docx")


def bold (docID, threshold=0.5):
    data = DB.get_json(docID)
    dic = addText.convertToText(data)
    document = Document()
    document.add_heading('Bold', 0)
    p= document.add_paragraph()

    for i in range(len(dic)):
        if (float(dic[i].get('wight')) < threshold):
            p.add_run(dic[i].get('content').replace("\n",' ')).bold = False

        else:
            p.add_run(dic[i].get('content').replace("\n",' ')).bold = True

    #document.add_page_break()
    document.save('demoBold.docx')
    convertDocxToImage("demoBold.docx")



def highligth(docID, color, threshold=0.5):

    data = DB.get_json(docID)
    dic = addText.convertToText(data)
    document = Document()
    document.add_heading('Highligth', 0)
    p = document.add_paragraph()

    for i in range(len(dic)):
        if (float(dic[i].get('wight')) < threshold):
            p.add_run(dic[i].get('content').replace("\n",' ') ).font

        else:
            font = p.add_run(dic[i].get('content').replace("\n",' ') ).font
            font.highlight_color = HighlightColor(color)

    #document.add_page_break()
    document.save('demoHighligth.docx')
    convertDocxToImage("demoHighligth.docx")



def HighlightColor(color):
    if (color.upper() == 'YELLOW'):
        return WD_COLOR_INDEX.YELLOW
    if (color.upper() == 'RED'):
        return WD_COLOR_INDEX.RED
    if (color.upper() == 'PINK'):
        return WD_COLOR_INDEX.PINK
    if (color.upper() == 'TURQUOISE'):  ##light blue
        return WD_COLOR_INDEX.TURQUOISE
    if (color.upper() == 'BRIGHT_GREEN'):
        return WD_COLOR_INDEX.BRIGHT_GREEN
    return WD_COLOR_INDEX.AUTO

def fontSize(docID, threshold=0.5):
    data = DB.get_json(docID)
    dic = addText.convertToText(data)
    document = Document()
    document.add_heading('Font Size', 0)
    p = document.add_paragraph()
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(12)
    for i in range(len(dic)):
        if (float(dic[i].get('wight')) < threshold):
            p.add_run(dic[i].get('content').replace("\n" , ' ')).font

        else:
            font = p.add_run(dic[i].get('content').replace("\n" , ' ')).font
            font.size = Pt(15)
    #document.add_page_break()
    document.save('demoFontSize.docx')
    convertDocxToImage("demoFontSize.docx")

def GradualFontSize(docID):
    lowRange=0
    midRange=0.3
    highRange=0.5
    higherRange=0.7

    data = DB.get_json(docID)
    dic = addText.convertToText(data)
    document = Document()
    document.add_heading('Gradual Font Size', 0)
    p = document.add_paragraph()
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(12)
    for i in range(len(dic)):
        threshold = float(dic[i].get('wight'))
        if (threshold >= lowRange and threshold <= midRange):
            font = p.add_run(dic[i].get('content').replace("\n", ' ')).font
            font.size = Pt(12)
        elif (threshold > midRange and threshold <= highRange):
            font = p.add_run(dic[i].get('content').replace("\n", ' ')).font
            font.size = Pt(14)
        elif (threshold > highRange and threshold <= higherRange):
            font = p.add_run(dic[i].get('content').replace("\n", ' ')).font
            font.size = Pt(16)
        else:
            font = p.add_run(dic[i].get('content').replace("\n", ' ')).font
            font.size = Pt(18)


    document.save('demoGradualFontSize.docx')
    convertDocxToImage("demoGradualFontSize.docx")


def fullText (docID):
    data = DB.get_json(docID)
    dic = addText.convertToText(data)
    document = Document()
    document.add_heading('Full Text', 0)
    p = document.add_paragraph()

    for i in range(len(dic)):
        p.add_run(dic[i].get('content').replace("\n",' '))
        #p.add_text(dic[i].get('content'))
    #document.add_page_break()
    document.save('demoFullText.docx')
    convertDocxToImage("demoFullText.docx")
    #document.add_page_break()
    document.save('fullText.docx')
    convertDocxToImage("fullText.docx")


def convertDocxToImage(fileName):
    # convert docx to pdf
    convert(fileName)
    # convert pdf to image
    newFileName = fileName.replace('.docx', '')
    images = convert_from_path(newFileName + '.pdf')
    for i, image in enumerate(images):
        fname = newFileName + str(i)+'.png'
        image.save(fname, "PNG")


def gradualHighlight(docID):
    data = DB.get_json(docID)
    dic = addText.convertToText(data)
    highlighted_text = []

    for i in range(len(dic)):
        highlighted_text.append(
                    '<span style="background-color:rgba(135,206,250,' + str(float(dic[i].get('wight'))) + ');">' +
                        dic[i].get('content') + '</span>')

    highlighted_text = ' '.join(highlighted_text)
    with open("demoGradualHighlight.html", "w") as htmlfile:
        htmlfile.write(highlighted_text)


    path_wkthmltoimage = r'C:\Program Files (x86)\wkhtmltopdf\bin\wkhtmltoimage.exe'
    config = imgkit.config(wkhtmltoimage=path_wkthmltoimage)

    imgkit.from_file('demoGradualHighlight.html', 'demoGradualHighlight.jpg', config=config)

